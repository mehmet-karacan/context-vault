"""Aşama 3.2: structural DOCX parser tests.

Builds a small ``.docx`` fixture programmatically (two headings, a paragraph,
a list, and a table interleaved to prove true body order) and asserts:

- paragraphs and tables are walked together in document body order;
- heading_path is correct on nested units;
- tables are produced as Markdown + structural JSON;
- list items are detected;
- block_index is monotonically increasing;
- no fabricated page numbers;
- ``NormalizedSource`` ``to_dict()`` / ``from_dict()`` round-trips losslessly.
"""

import io
import json

import pytest

from src.domain.normalized_content import NormalizedSource, UnitType
from src.infrastructure.parsers.docx_parser import DocxParser


def _build_fixture_bytes() -> bytes:
    """Builds a .docx in memory with a deliberate body order:
    H1 -> paragraph -> H2 -> list item -> table."""
    import docx
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    document = docx.Document()

    h1 = document.add_heading("Bölüm 1", level=1)
    document.add_paragraph("Paragraf 1")

    h2 = document.add_heading("Alt Konu", level=2)
    document.add_paragraph("Liste Paragrafı örneği", style="List Bullet")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Ad"
    table.cell(0, 1).text = "Değer"
    table.cell(1, 0).text = "PAYMENT_FLAG"
    table.cell(1, 1).text = "1"

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _parse(tmp_path) -> NormalizedSource:
    fixture = tmp_path / "fixture.docx"
    fixture.write_bytes(_build_fixture_bytes())
    return DocxParser().parse(str(fixture), "fixture.docx")


# --- parser contract -------------------------------------------------------


def test_supports_docx_extension_and_mime():
    parser = DocxParser()
    assert parser.supports("application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx")
    assert parser.supports("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "")
    assert parser.supports("", ".DOCX")
    assert not parser.supports("application/pdf", ".pdf")
    assert not parser.supports("text/plain", ".txt")


def test_source_metadata_and_shape(tmp_path):
    source = _parse(tmp_path)
    assert source.source_type == "document"
    assert source.title == "fixture.docx"
    assert source.language is None
    assert source.metadata["parser"] == "docx"
    assert "parser_profile" in source.metadata
    assert source.metadata["capabilities"]["textboxes"] is False
    assert "textboxes are not extracted" in source.metadata["capabilities"]["warnings"]


# --- body order ------------------------------------------------------------


def test_body_order_preserves_interleaved_paragraphs_and_tables(tmp_path):
    source = _parse(tmp_path)
    unit_types = [u.unit_type for u in source.units]

    headings = [i for i, t in enumerate(unit_types) if t == UnitType.HEADING]
    paragraphs = [i for i, t in enumerate(unit_types) if t == UnitType.PARAGRAPH]
    lists = [i for i, t in enumerate(unit_types) if t == UnitType.LIST_ITEM]
    tables = [i for i, t in enumerate(unit_types) if t == UnitType.TABLE]

    assert len(headings) == 2
    assert len(paragraphs) == 1
    assert len(lists) == 1
    assert len(tables) == 1

    assert headings[0] < paragraphs[0] < headings[1] < lists[0] < tables[0]


# --- heading path ----------------------------------------------------------


def test_heading_path_correct_on_nested_units(tmp_path):
    source = _parse(tmp_path)
    headings = [u for u in source.units if u.unit_type == UnitType.HEADING]
    h1, h2 = headings[0], headings[1]

    assert h1.hierarchy.heading_path == ["Bölüm 1"]
    assert h1.hierarchy.depth == 1

    para = [u for u in source.units if u.unit_type == UnitType.PARAGRAPH][0]
    assert para.hierarchy.heading_path == ["Bölüm 1"]

    assert h2.hierarchy.heading_path == ["Bölüm 1", "Alt Konu"]
    assert h2.hierarchy.depth == 2

    # Content under H2 carries the full ancestral heading path.
    list_item = [u for u in source.units if u.unit_type == UnitType.LIST_ITEM][0]
    assert list_item.hierarchy.heading_path == ["Bölüm 1", "Alt Konu"]

    table = [u for u in source.units if u.unit_type == UnitType.TABLE][0]
    assert table.hierarchy.heading_path == ["Bölüm 1", "Alt Konu"]


# --- tables ----------------------------------------------------------------
def test_table_produced_as_markdown_and_json(tmp_path):
    source = _parse(tmp_path)
    table = [u for u in source.units if u.unit_type == UnitType.TABLE][0]

    assert table.markdown.startswith("| Ad | Değer |")
    assert "|---|" in table.markdown
    assert "PAYMENT_FLAG" in table.markdown
    assert "1" in table.markdown

    assert table.metadata["rows"] == [
        ["Ad", "Değer"],
        ["PAYMENT_FLAG", "1"],
    ]
    assert table.metadata["row_count"] == 2
    assert table.metadata["col_count"] == 2


# --- list detection --------------------------------------------------------


def test_list_items_detected(tmp_path):
    source = _parse(tmp_path)
    assert any(u.unit_type == UnitType.LIST_ITEM for u in source.units)
    list_item = [u for u in source.units if u.unit_type == UnitType.LIST_ITEM][0]
    assert list_item.markdown.startswith("- ")
    assert list_item.metadata["list"] is True


# --- citation / block_index ------------------------------------------------


def test_block_index_monotonic_and_no_fabricated_pages(tmp_path):
    source = _parse(tmp_path)
    indexes = []
    for unit in source.units:
        assert unit.locator is not None
        assert unit.locator.block_index is not None
        assert unit.locator.page_start is None
        assert unit.locator.page_end is None
        assert unit.locator.bbox is None
        assert unit.locator.line_start is None
        assert unit.locator.line_end is None
        indexes.append(unit.locator.block_index)

    assert indexes == sorted(indexes)
    assert len(set(indexes)) == len(indexes)


# --- JSON round trip -------------------------------------------------------


def test_normalized_source_round_trips_losslessly(tmp_path):
    source = _parse(tmp_path)
    payload = json.dumps(source.to_dict(), ensure_ascii=False)
    restored = NormalizedSource.from_dict(json.loads(payload))
    assert restored == source
    assert restored.units == source.units
