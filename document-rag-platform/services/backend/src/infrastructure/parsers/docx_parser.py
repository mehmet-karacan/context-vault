"""Structural DOCX parser (Aşama 3.2).

Implements the ``DocumentParser`` port (``domain.ports``) for Office Open XML
Word documents and returns a shared ``NormalizedSource`` (Bölüm 6).

Behavior follows AKTIF_GOREV.md 3.2:

- Paragraphs and tables are walked together in true document body (XML)
  order by iterating ``doc.element.body`` children (``w:p`` and ``w:tbl``),
  never as two separate runs.
- Heading style levels (English ``Heading N`` / Turkish ``Başlık N``) and a
  heading-path stack are tracked so every emitted unit carries its ancestral
  heading context (heading_path).
- Tables are emitted as a ``ContentUnit(unit_type="table")``: a Markdown table
  in the ``markdown`` field plus the structural rows/cells in ``metadata``.
  Paragraph order inside a table cell is preserved.
- List / numbered paragraphs are detected best-effort (``numPr`` presence or a
  ``List*`` style) and emitted as ``unit_type="list_item"``.
- Empty-paragraph noise is cleaned, but section boundaries (headings, tables)
  are never dropped.
- Citation is ``heading_path`` + ``block_index``; no page numbers are
  fabricated (``page_start``/``page_end`` stay ``None`` for DOCX).
- Header/footer text is captured cheaply into source metadata; textboxes are
  not extracted and an explicit capability warning is recorded.
"""

from __future__ import annotations

import re
import uuid
from typing import List, Optional

from ...domain.normalized_content import (
    ContentUnit,
    Hierarchy,
    NormalizedSource,
    SourceLocator,
    UnitType,
)

_PARSER_NAME = "docx"
_PARSER_PROFILE = "struct-lite-v1"
_PARSER_VERSION = "0.1.0"

# Matches "Heading 1", "Heading 12", "Başlık 1", "Başlık 3", ...
_HEADING_RE = re.compile(r"^(?:heading|ba\u015fl\u0131k)\s*(\d+)$", re.IGNORECASE)


def _base_source(filename: Optional[str]) -> NormalizedSource:
    return NormalizedSource(
        source_id=str(uuid.uuid4()),
        source_type="document",
        title=filename,
        language=None,
    )


class DocxParser:
    """Structural parser for ``.docx`` files (Aşama 3.2)."""

    source_type = "docx"

    def supports(self, mime_type: str, extension: str) -> bool:
        ext = (extension or "").lower()
        if ext in (".docx",):
            return True
        mime = (mime_type or "").split(";", 1)[0].strip().lower()
        return (
            mime
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def parse(
        self,
        file_path: str,
        filename: str,
        options: Optional[dict] = None,
    ) -> NormalizedSource:
        import docx
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        options = options or {}
        document = docx.Document(file_path)

        source = _base_source(filename)
        source.metadata = {
            "parser": _PARSER_NAME,
            "parser_profile": _PARSER_PROFILE,
            "parser_version": _PARSER_VERSION,
            "parser_library": "python-docx " + (docx.__version__ or "unknown"),
            "origin": filename,
        }

        heading_stack: List[str] = []
        seq = 0

        for child in document.element.body.iterchildren():
            tag = child.tag
            if tag == qn("w:p"):
                seq += 1
                unit = self._paragraph_unit(
                    Paragraph(child, document),
                    filename=filename,
                    heading_stack=heading_stack,
                    seq=seq,
                )
                if unit is not None:
                    source.units.append(unit)
            elif tag == qn("w:tbl"):
                seq += 1
                unit = self._table_unit(
                    Table(child, document),
                    filename=filename,
                    heading_stack=heading_stack,
                    seq=seq,
                )
                if unit is not None:
                    source.units.append(unit)

        self._annotate_header_footer(source, document)
        return source

    # --- paragraph / heading / list --------------------------------------

    def _paragraph_unit(
        self,
        paragraph,
        filename: str,
        heading_stack: List[str],
        seq: int,
    ) -> Optional[ContentUnit]:
        text = self._paragraph_text(paragraph)
        style_name = self._style_name(paragraph)
        match = _HEADING_RE.match(style_name.strip())
        level = int(match.group(1)) if match else None

        if level is not None:
            heading_text = text if text else ""
            while len(heading_stack) >= level:
                heading_stack.pop()
            heading_stack.append(heading_text)

            unit = ContentUnit(
                unit_id=str(uuid.uuid4()),
                unit_type=UnitType.HEADING,
                text=heading_text,
                markdown=("#" * level + " " + heading_text).strip(),
                order=seq,
                hierarchy=Hierarchy(
                    heading_path=list(heading_stack),
                    parent_unit_id=None,
                    depth=level,
                ),
                locator=SourceLocator(file_path=filename, block_index=seq),
                metadata={"style": style_name, "heading_level": level},
            )
            return unit

        if not text:
            # Empty-paragraph noise: drop, but only when it is not a
            # section boundary (headings are handled above and never dropped).
            return None

        if self._is_list_paragraph(paragraph, style_name):
            unit_type = UnitType.LIST_ITEM
            markdown = "- " + text
            metadata = {"style": style_name, "list": True}
        else:
            unit_type = UnitType.PARAGRAPH
            markdown = text
            metadata = {"style": style_name}

        return ContentUnit(
            unit_id=str(uuid.uuid4()),
            unit_type=unit_type,
            text=text,
            markdown=markdown,
            order=seq,
            hierarchy=Hierarchy(
                heading_path=list(heading_stack),
                parent_unit_id=None,
                depth=len(heading_stack),
            ),
            locator=SourceLocator(file_path=filename, block_index=seq),
            metadata=metadata,
        )

    # --- tables -----------------------------------------------------------

    def _table_unit(
        self,
        table,
        filename: str,
        heading_stack: List[str],
        seq: int,
    ) -> ContentUnit:
        cells = self._table_cells(table)
        text = self._table_flat_text(cells)
        markdown = self._table_to_markdown(cells)

        return ContentUnit(
            unit_id=str(uuid.uuid4()),
            unit_type=UnitType.TABLE,
            text=text,
            markdown=markdown,
            order=seq,
            hierarchy=Hierarchy(
                heading_path=list(heading_stack),
                parent_unit_id=None,
                depth=len(heading_stack),
            ),
            locator=SourceLocator(file_path=filename, block_index=seq),
            metadata={
                "rows": cells,
                "row_count": len(cells),
                "col_count": max((len(r) for r in cells), default=0),
            },
        )

    @staticmethod
    def _table_cells(table) -> List[List[str]]:
        """Returns table as rows of cells; each cell preserves its internal
        paragraph order by joining them with newlines."""
        cells = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                paras = [p.text for p in cell.paragraphs if p.text.strip()]
                row_cells.append("\n".join(paras))
            cells.append(row_cells)
        return cells

    @staticmethod
    def _table_flat_text(cells: List[List[str]]) -> str:
        return "\n".join(" | ".join(cell for cell in row) for row in cells)

    @staticmethod
    def _table_to_markdown(cells: List[List[str]]) -> str:
        if not cells:
            return ""
        col_count = max(len(row) for row in cells)

        def esc(value: str) -> str:
            return value.replace("|", "\\|").replace("\n", "<br>")

        def fmt(row: List[str]) -> str:
            padded = row + [""] * (col_count - len(row))
            return "| " + " | ".join(esc(c) for c in padded) + " |"

        lines = [fmt(cells[0])]
        lines.append("|" + "|".join(["---"] * col_count) + "|")
        for row in cells[1:]:
            lines.append(fmt(row))
        return "\n".join(lines)

    # --- header / footer / textboxes -------------------------------------

    @staticmethod
    def _annotate_header_footer(source: NormalizedSource, document) -> None:
        header_texts: List[str] = []
        footer_texts: List[str] = []
        for section in document.sections:
            for header in (section.header, section.first_page_header,
                           section.even_page_header):
                try:
                    for p in header.paragraphs:
                        if p.text.strip():
                            header_texts.append(p.text.strip())
                except Exception:
                    continue
            for footer in (section.footer, section.first_page_footer,
                           section.even_page_footer):
                try:
                    for p in footer.paragraphs:
                        if p.text.strip():
                            footer_texts.append(p.text.strip())
                except Exception:
                    continue

        source.metadata["header_texts"] = header_texts
        source.metadata["footer_texts"] = footer_texts
        source.metadata["has_header_content"] = bool(header_texts)
        source.metadata["has_footer_content"] = bool(footer_texts)

        warnings = []
        if not bool(header_texts) and not bool(footer_texts):
            # No content found; still note support boundary explicitly.
            pass
        warnings.append("textboxes are not extracted")
        source.metadata["capabilities"] = {
            "body_order": True,
            "headings": True,
            "tables": True,
            "list_items": True,
            "header_footer": True,
            "textboxes": False,
            "warnings": warnings,
        }

    # --- small helpers ----------------------------------------------------

    @staticmethod
    def _paragraph_text(paragraph) -> str:
        return paragraph.text.strip()

    @staticmethod
    def _style_name(paragraph) -> str:
        try:
            return paragraph.style.name or ""
        except Exception:
            return ""

    @staticmethod
    def _is_list_paragraph(paragraph, style_name: str) -> bool:
        try:
            pPr = paragraph._p.pPr
            if pPr is not None and pPr.numPr is not None:
                return True
        except Exception:
            pass
        low = (style_name or "").lower()
        return "list" in low or low.startswith(("list", "liste", "madde", "numara"))
