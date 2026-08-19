import os
import re
import tempfile
import uuid
from datetime import datetime
from typing import List, Optional

from fastapi import Depends, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from sqlalchemy import func, or_
from sqlalchemy.orm import Session

from .db import get_db, init_db
from .llm import (
    AVAILABLE_CHAT_MODELS,
    CHAT_MODEL,
    PASSAGE_INSTRUCTION,
    QUERY_INSTRUCTION,
    embed_text,
    embed_texts,
    generate_answer,
)
from .models import Chunk, Document, Project

app = FastAPI(title="Document RAG API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
def on_startup():
    init_db()


class ChatQuery(BaseModel):
    query: str
    project_id: Optional[str] = None
    model: Optional[str] = None


class ProjectCreate(BaseModel):
    name: str


STOPWORDS = {
    "ve", "veya", "ile", "bir", "bu", "şu", "o", "de", "da", "mi", "mı", "mu", "mü",
    "nedir", "nasıl", "ne", "için", "gibi", "kadar", "çok", "az", "the", "and", "or",
    "is", "what", "how", "olan", "olarak", "var", "yok",
}


def extract_keywords(query: str) -> List[str]:
    words = re.findall(r"\w+", query.lower())
    return [w for w in words if len(w) >= 2 and w not in STOPWORDS]


# --- Text extraction -------------------------------------------------------


def extract_text_from_docx(file_path: str) -> str:
    import docx

    doc = docx.Document(file_path)
    return "\n".join(para.text for para in doc.paragraphs)


def extract_text_from_pdf(file_path: str) -> str:
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())


def extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def extract_text(file_path: str, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext in (".txt", ".md"):
        return extract_text_from_txt(file_path)
    raise ValueError(f"Desteklenmeyen dosya türü: {ext}")


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end]

        if end < text_length:
            last_period = chunk.rfind(".")
            last_newline = chunk.rfind("\n")
            break_point = max(last_period, last_newline)
            if break_point > chunk_size * 0.5:
                chunk = chunk[: break_point + 1]
                end = start + break_point + 1
            else:
                # No good sentence break — fall back to the last word
                # boundary so we never cut a word in half.
                last_space = chunk.rfind(" ")
                if last_space > 0:
                    chunk = chunk[:last_space]
                    end = start + last_space

        stripped = chunk.strip()
        if stripped:
            chunks.append(stripped)
        start = end - overlap

    return chunks


# --- Serialization -----------------------------------------------------------


def serialize_document(doc: Document, chunks_count: Optional[int] = None) -> dict:
    return {
        "id": str(doc.id),
        "name": doc.name,
        "size": doc.size,
        "status": doc.status,
        "uploaded_at": doc.uploaded_at.strftime("%Y-%m-%d %H:%M"),
        "chunks_count": chunks_count if chunks_count is not None else len(doc.chunks),
        "error_message": doc.error_message,
        "project_id": str(doc.project_id),
        "project_name": doc.project.name if doc.project else None,
    }


def serialize_project(project: Project, document_count: Optional[int] = None) -> dict:
    return {
        "id": str(project.id),
        "name": project.name,
        "created_at": project.created_at.strftime("%Y-%m-%d %H:%M"),
        "document_count": document_count if document_count is not None else len(project.documents),
    }


# --- Routes --------------------------------------------------------------


@app.get("/")
def root():
    return {"message": "Document RAG API is running"}


@app.get("/health")
def health(db: Session = Depends(get_db)):
    total = db.query(func.count(Document.id)).scalar()
    indexed = db.query(func.count(Document.id)).filter(Document.status == "indexed").scalar()
    return {
        "status": "healthy",
        "documents_count": total,
        "indexed_count": indexed,
    }


@app.get("/projects")
def list_projects(db: Session = Depends(get_db)):
    projects = db.query(Project).order_by(Project.created_at.desc()).all()
    return [serialize_project(p) for p in projects]


@app.post("/projects")
def create_project(payload: ProjectCreate, db: Session = Depends(get_db)):
    name = payload.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="Proje adı boş olamaz")
    if db.query(Project).filter(Project.name == name).first():
        raise HTTPException(status_code=409, detail="Bu isimde bir proje zaten var")
    project = Project(id=uuid.uuid4(), name=name, created_at=datetime.utcnow())
    db.add(project)
    db.commit()
    db.refresh(project)
    return serialize_project(project)


@app.delete("/projects/{project_id}")
def delete_project(project_id: str, db: Session = Depends(get_db)):
    project = db.get(Project, project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    db.delete(project)
    db.commit()
    return {"success": True, "message": f"Project {project_id} deleted"}


@app.get("/chat/models")
def list_chat_models():
    return {"models": AVAILABLE_CHAT_MODELS, "default": CHAT_MODEL}


@app.post("/documents/upload")
def upload_document(
    file: UploadFile = File(...),
    project_id: str = Form(...),
    chunk_size: int = Form(500),
    instruction: str = Form(PASSAGE_INSTRUCTION),
    db: Session = Depends(get_db),
):
    project = db.get(Project, project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    chunk_size = max(50, min(chunk_size, 4000))
    document = Document(
        id=uuid.uuid4(),
        project_id=project.id,
        name=file.filename,
        size=file.size or 0,
        status="uploaded",
        uploaded_at=datetime.utcnow(),
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    tmp_path = None
    try:
        suffix = os.path.splitext(file.filename)[1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(file.file.read())
            tmp_path = tmp_file.name

        document.status = "processing"
        db.commit()

        text = extract_text(tmp_path, file.filename)
        if not text.strip():
            raise ValueError("Belgenin okunabilir metin içeriği bulunamadı.")

        chunks = chunk_text(text, chunk_size=chunk_size, overlap=min(50, chunk_size // 5))
        embeddings = embed_texts(chunks, instruction=instruction)

        for index, (content, embedding) in enumerate(zip(chunks, embeddings)):
            db.add(
                Chunk(
                    id=uuid.uuid4(),
                    document_id=document.id,
                    chunk_index=index,
                    content=content,
                    embedding=embedding,
                )
            )

        document.status = "indexed"
        db.commit()
        db.refresh(document)

        return {
            "success": True,
            "document": serialize_document(document, chunks_count=len(chunks)),
            "message": f"{file.filename} {len(chunks)} parçaya bölünüp vektörlendi.",
        }

    except Exception as e:
        document.status = "error"
        document.error_message = str(e)
        db.commit()
        return {"success": False, "document": serialize_document(document), "error": str(e)}

    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


@app.get("/documents")
def list_documents(project_id: Optional[str] = None, db: Session = Depends(get_db)):
    query = db.query(Document)
    if project_id:
        query = query.filter(Document.project_id == project_id)
    documents = query.order_by(Document.uploaded_at.desc()).all()
    return [serialize_document(doc) for doc in documents]


@app.get("/documents/{doc_id}")
def get_document(doc_id: str, db: Session = Depends(get_db)):
    document = db.get(Document, doc_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return serialize_document(document)


@app.get("/documents/{doc_id}/status")
def get_document_status(doc_id: str, db: Session = Depends(get_db)):
    document = db.get(Document, doc_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    progress = {"uploaded": 20, "processing": 60, "indexed": 100, "error": 0}
    return {
        "id": str(document.id),
        "status": document.status,
        "chunks_count": len(document.chunks),
        "progress": progress.get(document.status, 0),
    }


@app.post("/documents/{doc_id}/delete")
def delete_document(doc_id: str, db: Session = Depends(get_db)):
    document = db.get(Document, doc_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    db.delete(document)
    db.commit()
    return {"success": True, "message": f"Document {doc_id} deleted"}


@app.post("/chat/query")
def query_chat(chat_query: ChatQuery, db: Session = Depends(get_db)):
    query_embedding = embed_text(chat_query.query, instruction=QUERY_INSTRUCTION)

    distance = Chunk.embedding.cosine_distance(query_embedding)
    search = (
        db.query(Chunk, Document, distance.label("distance"))
        .join(Document, Chunk.document_id == Document.id)
        .filter(Document.status == "indexed")
    )
    if chat_query.project_id:
        search = search.filter(Document.project_id == chat_query.project_id)

    TOP_K = 3

    # This is a per-corpus noise floor, not a fixed physical constant: chit-chat
    # queries ("selam") scored ~0.44-0.49 against one document set but up to
    # 0.53 against another (a relative top-vs-rest gap was tested and rejected
    # — it's *smaller* for genuine-but-weak matches than for "selam", so it
    # would reject real matches instead of chit-chat). 0.55 clears every
    # chit-chat probe measured so far; rely on the keyword-match path below
    # (not this threshold) to rescue literal term matches like "5G FWA" that
    # score low on pure vector similarity.
    SIMILARITY_THRESHOLD = 0.55
    vector_candidates = search.order_by(distance).limit(10).all()

    # Pure semantic ranking can bury a chunk that only mentions the query's
    # exact term once in passing (e.g. a product name cited in an unrelated
    # scoping paragraph) below chunks that are topically close but never
    # mention it at all — this is exactly what happened with "5G FWA": three
    # unrelated ETL chunks outranked the one chunk that names it. A chunk
    # containing a literal keyword from the query is pulled in regardless of
    # its vector rank, so an exact textual match is never lost to embedding
    # noise.
    keywords = extract_keywords(chat_query.query)
    keyword_candidates = []
    if keywords:
        keyword_search = search.filter(or_(*[Chunk.content.ilike(f"%{kw}%") for kw in keywords]))
        keyword_candidates = keyword_search.order_by(distance).limit(5).all()

    # Keyword matches get a reserved slot instead of competing with vector
    # candidates on similarity — that competition is exactly what buried
    # "5G FWA" in the first place. Fill remaining slots with the best
    # vector-only candidates above the threshold.
    keyword_ranked = sorted(
        [(chunk, doc, max(0.0, 1 - dist)) for chunk, doc, dist in keyword_candidates],
        key=lambda item: item[2],
        reverse=True,
    )[:TOP_K]
    seen_ids = {chunk.id for chunk, _, _ in keyword_ranked}

    vector_ranked = sorted(
        [
            (chunk, doc, max(0.0, 1 - dist))
            for chunk, doc, dist in vector_candidates
            if chunk.id not in seen_ids and max(0.0, 1 - dist) >= SIMILARITY_THRESHOLD
        ],
        key=lambda item: item[2],
        reverse=True,
    )

    ranked = keyword_ranked + vector_ranked[: TOP_K - len(keyword_ranked)]
    sources = [
        {"document": {"name": doc.name}, "chunk": chunk.content, "similarity": sim}
        for chunk, doc, sim in ranked
    ]

    answer = generate_answer(chat_query.query, [s["chunk"] for s in sources], model=chat_query.model)
    return {"answer": answer, "sources": sources}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
