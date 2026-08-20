"""Document upload, listing and lifecycle endpoints.

Moved verbatim from ``main.py`` — no behavior change. Text extraction and
chunking helpers live here because they are only used by the upload
endpoint; they will move to dedicated parser/chunker adapters in a later
stage (see AKTIF_GOREV.md Aşama 3/4).
"""

import hashlib
import mimetypes
import os
import tempfile
import uuid
from datetime import datetime
from typing import List, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from sqlalchemy.orm import Session

from ...config import settings
from ...db import get_db
from ...infrastructure.security import UploadValidationResult, validate_upload
from ...infrastructure.storage import object_keys
from ...infrastructure.storage.minio_storage import MinioObjectStorage
from ...llm import PASSAGE_INSTRUCTION, embed_texts
from ...models import Chunk, Document, DocumentVersion, IngestionJob, Project
from src.infrastructure.rate_limiter import rate_limiter

router = APIRouter(tags=["documents"])


def _build_object_storage() -> MinioObjectStorage:
    """Same construction the worker uses (Aşama 2.2's ``MinioObjectStorage``)
    — kept local to this module so the sync upload path (which never touched
    MinIO before Aşama 2.4) and the new async path share one code path."""
    return MinioObjectStorage(
        endpoint=settings.MINIO_ENDPOINT,
        access_key=settings.MINIO_ACCESS_KEY,
        secret_key=settings.MINIO_SECRET_KEY,
        bucket=settings.MINIO_BUCKET,
    )


# --- Text extraction -------------------------------------------------------


def extract_text_from_docx(file_path: str) -> str:
    import docx

    doc = docx.Document(file_path)
    return "\n".join(para.text for para in doc.paragraphs)


def extract_text_from_pdf(file_path: str) -> str:
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())


def extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def extract_text(file_path: str, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext in (".txt", ".md"):
        return extract_text_from_txt(file_path)
    raise ValueError(f"Desteklenmeyen dosya türü: {ext}")


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end]

        if end < text_length:
            last_period = chunk.rfind(".")
            last_newline = chunk.rfind("\n")
            break_point = max(last_period, last_newline)
            if break_point > chunk_size * 0.5:
                chunk = chunk[: break_point + 1]
                end = start + break_point + 1
            else:
                # No good sentence break — fall back to the last word
                # boundary so we never cut a word in half.
                last_space = chunk.rfind(" ")
                if last_space > 0:
                    chunk = chunk[:last_space]
                    end = start + last_space

        stripped = chunk.strip()
        if stripped:
            chunks.append(stripped)
        start = end - overlap

    return chunks


# --- Serialization -----------------------------------------------------------


def _latest_job_for_document(db: Session, document_id) -> Optional[IngestionJob]:
    """Returns the most recently created IngestionJob for any version of
    ``document_id``, or None if the document has never gone through the
    async ingestion pipeline (e.g. was uploaded before Aşama 2.4, or was
    uploaded with ``FEATURE_ASYNC_INGESTION=False``).

    Additive lookup only — never used to decide anything about the document
    row itself, just to surface job progress alongside it (Aşama 2.4 kabul
    kriteri: existing fields are never removed or repurposed).
    """
    return (
        db.query(IngestionJob)
        .join(DocumentVersion, IngestionJob.version_id == DocumentVersion.id)
        .filter(DocumentVersion.document_id == document_id)
        .order_by(IngestionJob.created_at.desc())
        .first()
    )


def serialize_document(
    doc: Document,
    chunks_count: Optional[int] = None,
    job: Optional[IngestionJob] = None,
) -> dict:
    return {
        "id": str(doc.id),
        "name": doc.name,
        "size": doc.size,
        "status": doc.status,
        "uploaded_at": doc.uploaded_at.strftime("%Y-%m-%d %H:%M"),
        "chunks_count": chunks_count if chunks_count is not None else len(doc.chunks),
        "error_message": doc.error_message,
        "project_id": str(doc.project_id),
        "project_name": doc.project.name if doc.project else None,
        # --- Aşama 2.4 additive fields: only populated when this document
        # has an associated IngestionJob (async pipeline). None for
        # documents ingested synchronously / before this stage.
        "job_id": str(job.id) if job is not None else None,
        "job_status": job.status if job is not None else None,
        "job_stage": job.stage if job is not None else None,
        "job_error": job.error_message if job is not None else None,
    }


# --- Security ----------------------------------------------------------------


def guard_upload(file_bytes: bytes, filename: str, mime_type: str) -> None:
    """Reject an upload that fails MIME/magic/size/safety validation.

    Aşama 9.5 guard: on violation we raise an HTTP 400 with a short, safe
    message — never a stack trace. Non-breaking: only rejects clearly unsafe
    inputs (oversize, extension-vs-magic mismatch, total-limit breach).
    """
    result: UploadValidationResult = validate_upload(file_bytes, filename, mime_type)
    if not result.ok:
        raise HTTPException(status_code=400, detail=result.error or "Upload rejected")


# --- Routes --------------------------------------------------------------


def _upload_document_async(
    file: UploadFile,
    project: Project,
    db: Session,
) -> dict:
    """FEATURE_ASYNC_INGESTION=True path (Aşama 2.4): persists the Document,
    a first DocumentVersion and a queued IngestionJob immediately, writes the
    original bytes to MinIO, enqueues the Celery task and returns without
    waiting for parse/chunk/embed. The document is visible in ``GET
    /documents`` the instant this function returns, with
    ``status="uploaded"`` moving to ``"processing"``/``"indexed"``/``"error"``
    as the worker (``workers.ingestion_tasks.run_ingestion_job``) advances.
    """
    # Local import: ``ingestion_tasks`` imports ``chunk_text``/``extract_text``
    # from this module at module load time, so importing it back at module
    # scope here would create a circular import. Deferring to call time is
    # safe because both modules are fully loaded by the time a request
    # actually arrives.
    from ...workers.ingestion_tasks import process_ingestion_job

    file_bytes = file.file.read()
    guard_upload(file_bytes, file.filename or "", file.content_type or "")
    document_id = uuid.uuid4()
    version_id = uuid.uuid4()
    job_id = uuid.uuid4()
    now = datetime.utcnow()

    document = Document(
        id=document_id,
        project_id=project.id,
        name=file.filename,
        size=len(file_bytes) if not file.size else file.size,
        status="uploaded",
        uploaded_at=now,
        source_type="document",
        mime_type=file.content_type or mimetypes.guess_type(file.filename or "")[0],
        checksum=hashlib.sha256(file_bytes).hexdigest(),
        created_at=now,
        updated_at=now,
    )
    db.add(document)

    # New document -> this is always its first version (Aşama 2 kabul
    # kriteri: re-index/version_no computation only matters once a document
    # can be re-ingested, which is out of this endpoint's scope — see
    # AKTIF_GOREV.md Aşama 2.4 task notes).
    existing_max_version_no = (
        db.query(DocumentVersion)
        .filter(DocumentVersion.document_id == document_id)
        .count()
    )
    version = DocumentVersion(
        id=version_id,
        document_id=document_id,
        version_no=existing_max_version_no + 1,
        status="pending",
        created_at=now,
    )
    db.add(version)

    storage = _build_object_storage()
    storage_key = object_keys.original_key(
        str(project.id), str(document_id), str(version_id), file.filename or "file"
    )
    storage.put(storage_key, file_bytes, content_type=file.content_type)
    version.storage_key = storage_key

    job = IngestionJob(
        id=job_id,
        version_id=version_id,
        status="queued",
        stage="validating",
        attempt=0,
        created_at=now,
    )
    db.add(job)

    db.commit()
    db.refresh(document)

    # Enqueue after commit: if the commit itself fails, no orphaned job is
    # ever picked up by a worker.
    process_ingestion_job.delay(str(job_id))

    return {
        "job_id": str(job_id),
        "document_id": str(document_id),
        "version_id": str(version_id),
        "status": "queued",
        "document": serialize_document(document, chunks_count=0, job=job),
    }


@router.post("/documents/upload")
def upload_document(
    file: UploadFile = File(...),
    project_id: str = Form(...),
    chunk_size: int = Form(500),
    instruction: str = Form(PASSAGE_INSTRUCTION),
    _: None = Depends(rate_limiter),
    db: Session = Depends(get_db),
):
    project = db.get(Project, project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    if settings.FEATURE_ASYNC_INGESTION:
        return _upload_document_async(file, project, db)

    chunk_size = max(50, min(chunk_size, 4000))
    document = Document(
        id=uuid.uuid4(),
        project_id=project.id,
        name=file.filename,
        size=file.size or 0,
        status="uploaded",
        uploaded_at=datetime.utcnow(),
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    tmp_path = None
    try:
        file_bytes = file.file.read()
        guard_upload(file_bytes, file.filename or "", file.content_type or "")
        suffix = os.path.splitext(file.filename)[1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name

        document.status = "processing"
        db.commit()

        text = extract_text(tmp_path, file.filename)
        if not text.strip():
            raise ValueError("Belgenin okunabilir metin içeriği bulunamadı.")

        chunks = chunk_text(text, chunk_size=chunk_size, overlap=min(50, chunk_size // 5))
        embeddings = embed_texts(chunks, instruction=instruction)

        for index, (content, embedding) in enumerate(zip(chunks, embeddings)):
            db.add(
                Chunk(
                    id=uuid.uuid4(),
                    document_id=document.id,
                    chunk_index=index,
                    content=content,
                    embedding=embedding,
                )
            )

        document.status = "indexed"
        db.commit()
        db.refresh(document)

        return {
            "success": True,
            "document": serialize_document(document, chunks_count=len(chunks)),
            "message": f"{file.filename} {len(chunks)} parçaya bölünüp vektörlendi.",
        }

    except Exception as e:
        document.status = "error"
        document.error_message = str(e)
        db.commit()
        return {"success": False, "document": serialize_document(document), "error": str(e)}

    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


@router.get("/documents")
def list_documents(project_id: Optional[str] = None, db: Session = Depends(get_db)):
    query = db.query(Document)
    if project_id:
        query = query.filter(Document.project_id == project_id)
    documents = query.order_by(Document.uploaded_at.desc()).all()
    return [
        serialize_document(doc, job=_latest_job_for_document(db, doc.id)) for doc in documents
    ]


@router.get("/documents/{doc_id}")
def get_document(doc_id: str, db: Session = Depends(get_db)):
    document = db.get(Document, doc_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return serialize_document(document, job=_latest_job_for_document(db, document.id))


@router.get("/documents/{doc_id}/status")
def get_document_status(doc_id: str, db: Session = Depends(get_db)):
    document = db.get(Document, doc_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    progress = {"uploaded": 20, "processing": 60, "indexed": 100, "error": 0}
    return {
        "id": str(document.id),
        "status": document.status,
        "chunks_count": len(document.chunks),
        "progress": progress.get(document.status, 0),
    }


@router.post("/documents/{doc_id}/delete")
def delete_document(doc_id: str, db: Session = Depends(get_db)):
    document = db.get(Document, doc_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    db.delete(document)
    db.commit()
    return {"success": True, "message": f"Document {doc_id} deleted"}
